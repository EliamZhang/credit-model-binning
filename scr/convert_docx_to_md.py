"""将 分箱方法论说明.docx 转换为 Markdown。

一次性转换工具：按文档元素顺序输出标题、段落、列表与表格；
合并单元格行（如候选评分表的脚注行）抽取为表后斜体说明。
"""
import sys

import docx
from docx.table import Table
from docx.text.paragraph import Paragraph

SRC = "分箱方法论说明.docx"
DST = "分箱方法论说明.md"

HEADING_LEVELS = {
    "Heading 1": 1,
    "Heading 2": 2,
    "Heading 3": 3,
    "标题 1": 1,
    "标题 2": 2,
    "标题 3": 3,
}


def cell_text(cell) -> str:
    texts = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
    return "\n".join(texts)


def para_is_bold(p: Paragraph) -> bool:
    runs = [r for r in p.runs if r.text.strip()]
    return bool(runs) and all(r.bold for r in runs)


def render_table(t: Table) -> str:
    ncols = len(t.columns)
    rows: list[list[str]] = []
    footnotes: list[str] = []

    for row in t.rows:
        seen: set[int] = set()
        cells: list[str] = []
        merged = False
        for c in row.cells:
            tc_id = id(c._tc)
            if tc_id in seen:
                merged = True
                continue
            seen.add(tc_id)
            cells.append(cell_text(c))
        if merged:
            footnotes.append("\n".join(cells))
        else:
            rows.append(cells)

    lines = []
    for ri, cells in enumerate(rows):
        cells = cells + [""] * (ncols - len(cells))
        escaped = [c.replace("|", "\\|").replace("\n", "<br>") for c in cells]
        lines.append("| " + " | ".join(escaped) + " |")
        if ri == 0:
            lines.append("|" + "---|" * ncols)
    for fn in footnotes:
        lines.append("")
        lines.append("*" + fn.replace("|", "\\|").replace("\n", " ") + "*")
    return "\n".join(lines)


def render_flowchart(t: Table) -> str:
    lines = []
    for row in t.rows:
        for c in row.cells:
            for p in c.paragraphs:
                if p.text.strip():
                    lines.append(p.text.strip())
    return "\n".join("> " + ln for ln in lines)


def main() -> None:
    doc = docx.Document(SRC)
    body = doc.element.body

    blocks: list[str] = []
    seen_first_heading = False

    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]

        if tag == "p":
            p = Paragraph(child, doc)
            text = p.text.strip().replace("\t", " ")
            if not text:
                continue
            style = p.style.name if p.style else ""
            level = HEADING_LEVELS.get(style)

            if level:
                blocks.append(f"{'#' * level} {text}")
                seen_first_heading = True
            elif not seen_first_heading:
                # 标题页与目录区：按字号识别大标题，其余原样保留。
                big = any(
                    r.font.size and r.font.size.pt >= 20 and r.text.strip()
                    for r in p.runs
                )
                if big:
                    blocks.append(f"# {text}")
                elif text.startswith("目"):
                    blocks.append(f"## {text}")
                elif text.startswith("评审稿"):
                    blocks.append(f"*{text}*")
                else:
                    blocks.append(text)
            elif text.startswith("•"):
                blocks.append("- " + text.lstrip("•").strip())
            else:
                prefix = "**" if para_is_bold(p) else ""
                suffix = "**" if prefix else ""
                blocks.append(prefix + text + suffix)

        elif tag == "tbl":
            table = Table(child, doc)
            if table.rows[0].cells[0].text.strip().startswith("样本切分"):
                blocks.append(render_flowchart(table))
            else:
                blocks.append(render_table(table))

    # 合并相邻列表项，避免松散列表的额外空行。
    merged: list[str] = []
    for b in blocks:
        if b.startswith("- ") and merged and merged[-1].startswith("- "):
            merged[-1] += "\n" + b
        else:
            merged.append(b)

    for i, b in enumerate(merged):
        if b.lstrip().startswith("# 一、"):
            merged.insert(i, "---")
            break

    with open(DST, "w", encoding="utf-8") as f:
        f.write("\n\n".join(merged).strip() + "\n")
    print(f"written: {DST}")


if __name__ == "__main__":
    sys.exit(main())
