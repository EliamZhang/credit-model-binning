# -*- coding: utf-8 -*-
"""生成《模型分数分箱与策略阈值设定方法说明》商业风格（美化版）Word 文档。

排版规范：
- 统一深蓝商务配色：主色 #1F4E79，斑马纹 #EAF1F8，提示行 #FFF8E6
- 第 1 页标题块 + 目录同页：顶部装饰线右置“评审稿 V1.0”、加宽字距大标题，无页眉页脚
- 正文从第 2 页起：页眉左文档名 / 右“评审稿 V1.0”，主题色分隔线；页脚灰色“第 X 页 共 Y 页”
- 一级标题黑体四号深蓝 + 底部横线并分页；二级标题黑体小四深蓝
- 表格：深蓝表头白字、隔行浅蓝斑马纹、首列加粗、浅灰细边框；表格末行可附浅黄提示合并行
- 优先级分组标签：浅蓝底纹黑体深蓝加粗（用于"局限性与待讨论问题"三级分级）
- 正文宋体小四、1.5 倍行距、首行缩进 2 字符、两端对齐
- 公式居中、编号右对齐；表题居表上、图题居图下（黑体五号深蓝）
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

EAST_FONT = "宋体"
HEI_FONT = "黑体"
KAI_FONT = "楷体"
BODY_SIZE = 12          # 小四
SMALL_SIZE = 10.5       # 五号
PRIMARY = "1F4E79"      # 深蓝（主色）
BAND_FILL = "EAF1F8"    # 浅蓝（斑马纹 / 流程框）
NOTE_FILL = "FFF8E6"    # 浅黄（表末提示行）
TABLE_BORDER = "BFBFBF"  # 浅灰（表格边框）
GRAY = "595959"          # 灰（页脚辅助文字）
CONTENT_WIDTH_CM = 15.5  # 21 - 3.0 - 2.5
DOC_TITLE = "模型分数分箱与策略阈值设定方法说明"
OUTPUT = "分箱方法论说明.docx"


def rgb(hex_str):
    return RGBColor.from_string(hex_str)


def set_run_font(run, ea_font=EAST_FONT, size=BODY_SIZE, bold=False, color=None):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), ea_font)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_char_spacing(run, val):
    """设置字符间距（1/20 pt）。"""
    rPr = run._element.get_or_add_rPr()
    sp = rPr.find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing")
        rPr.append(sp)
    sp.set(qn("w:val"), str(val))


def setup_section(section, left=3.0, right=2.5):
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


def set_pg_num_type(section, fmt=None, start=1):
    sect_pr = section._sectPr
    for old in sect_pr.findall(qn("w:pgNumType")):
        sect_pr.remove(old)
    pg = OxmlElement("w:pgNumType")
    if fmt:
        pg.set(qn("w:fmt"), fmt)
    pg.set(qn("w:start"), str(start))
    sect_pr.append(pg)


def setup_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_FONT)
    normal.font.size = Pt(BODY_SIZE)
    for level, size in [(1, 14), (2, 12)]:
        st = doc.styles[f"Heading {level}"]
        st.font.name = "Times New Roman"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), HEI_FONT)
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = rgb(PRIMARY)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table, color=TABLE_BORDER, sz="4"):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:color"), color)
        borders.append(el)
    tbl_pr.append(borders)


def add_para_border(p, edge="bottom", color=PRIMARY, sz="8", space="4"):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), sz)
    el.set(qn("w:space"), space)
    el.set(qn("w:color"), color)
    pbdr.append(el)
    spacing = pPr.find(qn("w:spacing"))
    if spacing is not None:
        spacing.addprevious(pbdr)
    else:
        pPr.append(pbdr)


def add_field_run(p, instr, size=9, ea_font=EAST_FONT, color=None, bold=False):
    r = p.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
    r._r.append(f1); r._r.append(it); r._r.append(f2)
    set_run_font(r, ea_font, size, bold=bold, color=color)
    return r


def blank(doc, size=BODY_SIZE, count=1):
    for _ in range(count):
        p = doc.add_paragraph()
        set_run_font(p.add_run(""), EAST_FONT, size)


def hrule(doc, color=PRIMARY, sz="16", space="1"):
    """装饰横线（空段落 + 底边框）。"""
    p = doc.add_paragraph()
    add_para_border(p, "bottom", color, sz, space)
    return p


def centered(doc, text, ea_font=HEI_FONT, size=BODY_SIZE, bold=False, space_after=6, color=None, char_spacing=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_run_font(r, ea_font, size, bold=bold, color=color)
    if char_spacing:
        set_char_spacing(r, char_spacing)
    return p


def body(doc, text, indent=True, bold=False, size=BODY_SIZE):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_run_font(p.add_run(text), EAST_FONT, size, bold=bold)
    return p


def bullet(doc, text, size=BODY_SIZE, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(24)
    p.paragraph_format.first_line_indent = Pt(-12)
    p.paragraph_format.line_spacing = 1.4
    p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix:
        set_run_font(p.add_run("• " + bold_prefix), EAST_FONT, size, bold=True, color=rgb(PRIMARY))
        set_run_font(p.add_run(text), EAST_FONT, size)
    else:
        set_run_font(p.add_run("• " + text), EAST_FONT, size)
    return p


def priority_label(doc, text):
    """优先级分组标签：浅蓝底纹黑体深蓝加粗，用于"局限性与待讨论问题"分级。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), BAND_FILL)
    pPr.append(shd)
    set_run_font(p.add_run(text), HEI_FONT, 11, bold=True, color=rgb(PRIMARY))
    return p


def formula(doc, text, number=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    if number:
        p.paragraph_format.tab_stops.add_tab_stop(
            Cm(CONTENT_WIDTH_CM), WD_TAB_ALIGNMENT.RIGHT
        )
        text = f"{text}\t({number})"
    set_run_font(p.add_run(text), "Times New Roman", 12)
    return p


def heading(doc, level, text):
    p = doc.add_heading(level=level)
    p.add_run(text)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    if level == 1:
        p.paragraph_format.page_break_before = True
        add_para_border(p, "bottom", PRIMARY, sz="12", space="4")
    return p


def caption(doc, text, below=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4 if not below else 8)
    p.paragraph_format.space_before = Pt(8 if not below else 4)
    p.paragraph_format.keep_with_next = True
    set_run_font(p.add_run(text), HEI_FONT, SMALL_SIZE, bold=True, color=rgb(PRIMARY))
    return p


def add_table(doc, headers, rows, widths=None, caption_text=None, footnote_lines=None):
    """footnote_lines：附加到表末的浅黄提示合并行（跨全列、宋体 11pt）。"""
    if caption_text:
        caption(doc, caption_text)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color=TABLE_BORDER, sz="4")
    hdr_row = table.rows[0]
    tr_pr = hdr_row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(cell, PRIMARY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(h), HEI_FONT, SMALL_SIZE, bold=True, color=rgb("FFFFFF"))
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, v in enumerate(row):
            cell = cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_idx % 2 == 1:
                shade_cell(cell, BAND_FILL)
            p = cell.paragraphs[0]
            if i == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_run_font(p.add_run(str(v)), HEI_FONT, SMALL_SIZE, bold=True)
            else:
                set_run_font(p.add_run(str(v)), EAST_FONT, SMALL_SIZE)
    if widths:
        for i, w in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    if footnote_lines:
        merged = table.rows[-1].cells[0].merge(table.rows[-1].cells[-1])
        shade_cell(merged, NOTE_FILL)
        first_p = merged.paragraphs[0]
        for extra in merged.paragraphs[1:]:
            merged._tc.remove(extra._p)
        first_p.clear()
        for i, line in enumerate(footnote_lines):
            p = first_p if i == 0 else merged.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing = 1.4
            p.paragraph_format.space_after = Pt(2)
            set_run_font(p.add_run(line), EAST_FONT, 11)
    return table


def flow_box(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color=PRIMARY, sz="12")
    cell = table.rows[0].cells[0]
    shade_cell(cell, BAND_FILL)
    cell.width = Cm(13.5)
    for i, line in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(line)
        if i == 0:
            set_run_font(r, HEI_FONT, 11, bold=True, color=rgb(PRIMARY))
        else:
            set_run_font(r, EAST_FONT, 11)
    return table


def add_toc(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
    it.text = 'TOC \\o "1-2" \\h \\z \\u'
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "（目录域：在 Word 中右键此处选择“更新域”，或 Ctrl+A 后按 F9 生成目录）"
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end")
    r._r.append(f1); r._r.append(it); r._r.append(f2); r._r.append(t); r._r.append(f3)
    set_run_font(r, EAST_FONT, BODY_SIZE)


def add_body_header_footer(section):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.paragraph_format.tab_stops.add_tab_stop(Cm(CONTENT_WIDTH_CM), WD_TAB_ALIGNMENT.RIGHT)
    r1 = p.add_run(DOC_TITLE)
    set_run_font(r1, EAST_FONT, 9, color=rgb(GRAY))
    r2 = p.add_run("\t评审稿 V1.0")
    set_run_font(r2, KAI_FONT, 9, bold=True, color=rgb(PRIMARY))
    add_para_border(p, "bottom", PRIMARY, sz="6", space="1")

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(fp.add_run("第 "), EAST_FONT, 9, color=rgb(GRAY))
    add_field_run(fp, "PAGE", size=9, color=rgb(GRAY))
    set_run_font(fp.add_run(" 页 共 "), EAST_FONT, 9, color=rgb(GRAY))
    add_field_run(fp, "NUMPAGES", size=9, color=rgb(GRAY))
    set_run_font(fp.add_run(" 页"), EAST_FONT, 9, color=rgb(GRAY))


def build_title_page(doc):
    """第 1 页标题块：顶部装饰线（右置“评审稿 V1.0”）+ 项目名 + 大标题 + 副标题。"""
    p = doc.add_paragraph()
    p.paragraph_format.tab_stops.add_tab_stop(Cm(CONTENT_WIDTH_CM), WD_TAB_ALIGNMENT.RIGHT)
    set_run_font(p.add_run("\t评审稿 · V1.0"), KAI_FONT, 10.5, bold=True, color=rgb(GRAY))
    add_para_border(p, "bottom", PRIMARY, sz="12", space="6")
    blank(doc, count=2)

    centered(doc, "项目：消费信贷评分模型分箱与策略优化", KAI_FONT, 12, color=rgb(GRAY), space_after=2)
    blank(doc, count=2)
    centered(doc, "模型分数分箱与策略阈值设定方法说明", HEI_FONT, 26, bold=True, color=rgb(PRIMARY), space_after=4, char_spacing=40)
    centered(doc, "—— 等频初分 · 约束合箱 · 候选择优 · 阈值决策 · 独立验证 ——", KAI_FONT, 12, color=rgb(GRAY), space_after=2, char_spacing=20)
    blank(doc, count=2)


def main():
    doc = Document()
    setup_section(doc.sections[0])
    setup_styles(doc)
    doc.core_properties.title = DOC_TITLE

    # ============ 第 1 节：标题块 + 目录（无页眉页脚、无页码） ============
    build_title_page(doc)
    centered(doc, "目　　录", HEI_FONT, 15, bold=True, color=rgb(PRIMARY), space_after=4, char_spacing=60)
    hrule(doc, PRIMARY, sz="12", space="4")
    add_toc(doc)

    # ============ 第 2 节：正文（页眉页脚，页码从 1 起） ============
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    setup_section(body_section)
    add_body_header_footer(body_section)
    set_pg_num_type(body_section, start=1)

    upd = OxmlElement("w:updateFields")
    upd.set(qn("w:val"), "true")
    doc.settings.element.append(upd)

    # ============ 一、总体思路 ============
    heading(doc, 1, "一、总体思路")
    body(doc, "消费信贷审批中，模型输出连续风险分（score_mlt），业务侧需要将其转化为稳定、单调、可解释的风险等级，并据此确定自动通过、人工审核、拒绝三段阈值。本文档说明当前分箱与阈值方案的方法论与验证逻辑，供评审确认。")
    body(doc, "整体流程：在完整 Train 上等频初分 → 约束自动合箱为 6~8 档（目标 7 档）→ 候选方案评分择优 → 映射至 OOT → 风险约束下确定阈值 → OOT 独立验证 → 输出 Excel 策略报告。")
    flow_box(doc, [
        "样本切分（Train / OOT）",
        "→ Train 上 20 等频初分 → 边界复用到 OOT",
        "→ 自动合箱（约束修正 → 单调合并 → 档位压缩 → 候选生成）",
        "→ Train 上候选方案评分与选择",
        "→ 验证（单调性 / PSI / AUC-KS / 月度稳定性）",
        "→ 阈值曲线与三段策略 → 输出报告",
    ])
    caption(doc, "图 1　整体技术路线", below=True)

    # ============ 二、样本设计与风险指标 ============
    heading(doc, 1, "二、样本设计与风险指标")

    heading(doc, 2, "（一）样本切分")
    add_table(
        doc,
        ["数据集", "时间范围", "职责"],
        [
            ["Train", "截止月份及以前", "学习分箱边界、执行合箱、选择方案并确定策略阈值"],
            ["OOT（时间外验证）", "Train 截止月份之后", "独立验证，不参与任何合箱与阈值调参"],
        ],
        widths=[3.2, 5.6, 7.2],
        caption_text="表 1　样本切分设计",
    )
    bullet(doc, "分箱边界仅在 Train 上学习，并原样复用到 OOT，保证两个数据集的分箱口径一致。")
    bullet(doc, "OOT 全程不参与方案开发，仅用于最终样本外验证。")

    heading(doc, 2, "（二）风险指标")
    add_table(
        doc,
        ["维度", "指标", "定义", "作用"],
        [
            ["观察窗口", "1M30+", "到期 1 个月内逾期 30 天以上", "短期风险，反应快"],
            ["观察窗口", "3M30+", "到期 3 个月内逾期 30 天以上", "中期风险，更成熟稳定"],
            ["计量口径", "笔数逾期率", "逾期样本数 ÷ 成熟样本数", "人群风险暴露面"],
            ["计量口径", "金额逾期率", "逾期剩余本金 ÷ 成熟本金敞口", "金额损失强度"],
        ],
        widths=[2.4, 2.6, 6.4, 4.6],
        caption_text="表 2　风险指标体系",
    )
    bullet(doc, "成熟样本定义：表现标签已确定（0/1）的样本才进入逾期率分母，未成熟样本不参与。")
    bullet(doc, "合箱主指标：1M30+ 与 3M30+ 笔数逾期率双指标，任一出现相邻倒挂即触发合并；显著性检验锚定 3M30+。金额口径用于候选评价与策略验证。")

    # ============ 三、分箱方法论 ============
    heading(doc, 1, "三、分箱方法论")

    heading(doc, 2, "（一）初分与档位目标")
    body(doc, "在完整 Train 上按分数分位数做 20 等频初分，区间规则 (left, right]，首尾扩展为 ±∞；分数重复过多时合并相同边界，实际箱数可能少于 20。自动合箱至 6~8 档，目标 7 档。")

    heading(doc, 2, "（二）单箱硬约束")
    add_table(
        doc,
        ["约束项", "普通箱要求", "说明"],
        [
            ["样本占比", "中间箱 ≥ 5%，头尾箱 ≥ 2.5%", "尾部箱允许更小，保证极端风险可分档"],
            ["主指标成熟量", "≥ 1,000", "保证逾期率估计的统计稳健性"],
            ["坏样本量", "≥ 20", "保证风险估计不是由个别样本驱动"],
            ["好样本量", "≥ 200", "保证低风险估计的稳定性"],
        ],
        widths=[3.4, 5.6, 7.0],
        caption_text="表 3　最终档位单箱硬约束",
    )
    body(doc, "同时引入极端人群圈选保护：默认把最好（低风险端）与最坏（高风险端）各 1 个初始箱圈定为极端箱，不允许常规合箱跨越其边界，并放宽约束（成熟量 ≥ 500，最好箱坏样本下限 0、最坏箱好样本下限 0）。最好 / 最坏人群分别对应自动通过与拒绝策略的关键结构，不能被普通合箱稀释。")

    heading(doc, 2, "（三）单调性")
    bullet(doc, "Train：主指标（1M30+、3M30+ 笔数逾期率）相邻箱不允许倒挂（容差 0）。")
    bullet(doc, "候选评分同时监控 Train 上四类风险率（含金额口径）的倒挂数。")
    bullet(doc, "月度稳定性检查允许 0.3 个百分点的容忍倒挂。")
    body(doc, "单调性是风险等级可解释、可监控的基础，也是业务沟通与合规评审的基本要求。")

    heading(doc, 2, "（四）保护边界")
    body(doc, "以下关键边界在自动合箱中尽量保留，合并跨越需付出高额代价（见（六））：")
    bullet(doc, "策略风险边界：自动通过与总接纳约束对应的累计 3M30+ 风险线与边际风险线位置。")
    bullet(doc, "风险跳升边界：相邻箱之间 3M30+ 风险跳升最大的边界（默认保护 1 个）。")
    bullet(doc, "极端人群圈选边界：最好 / 最坏极端箱的外侧边界（默认硬保护，禁止跨越）。")
    body(doc, "策略阈值最终落在箱边界上，保护边界防止合箱吞掉策略关键结构。")

    heading(doc, 2, "（五）合箱流程")
    add_table(
        doc,
        ["阶段", "触发条件", "合并策略"],
        [
            ["① 约束修正", "存在违反单箱硬约束的箱", "优先处理违反最严重的箱，仅允许与其相邻箱合并"],
            ["② 单调合并（PAVA 风格）", "主指标存在相邻倒挂", "从倒挂最严重的一对开始合并，直至无倒挂"],
            ["③ 档位压缩", "档位数仍多于 8 档", "反复合并合并代价最低的相邻对"],
            ["④ 候选生成", "档位数仍多于 6 档", "优先合并“风险差异不显著（p ≥ 0.10）或差距 ≤ 0.3%”的相邻对，生成 8、7、6 档候选"],
        ],
        widths=[2.8, 5.2, 8.0],
        caption_text="表 4　合箱四阶段",
    )
    body(doc, "每个阶段记录合并原因与前后状态，全过程可追溯。")

    heading(doc, 2, "（六）合并代价")
    formula(doc, "merge_cost = 风险率差距 × 100 + (1 − 两比例Z检验p值) + IV损失 × 10 + 保护边界惩罚", "1")
    formula(doc, "保护边界惩罚：普通保护边界 100；极端圈选边界 10,000", "2")
    body(doc, "风险差距越大、差异越显著、IV 损失越大、跨越保护边界，代价越高，越不应合并。")

    heading(doc, 2, "（七）候选评分与选择")
    body(doc, "四阶段产生的每一步合箱状态均为候选方案，先做硬约束判定，再对可行方案打分择优：")
    bullet(doc, "硬约束（必须全部满足）：档位数 6~8；Train 主指标无倒挂；单箱约束全部满足；未跨越极端圈选边界。")
    add_table(
        doc,
        ["评分项", "权重", "含义"],
        [
            ["硬约束全部满足", "+100", "可行方案的基础分"],
            ["Train 主指标倒挂", "−30 × 次数", "惩罚主指标单调性缺陷"],
            ["Train 全指标倒挂", "−4 × 次数", "惩罚辅助指标倒挂"],
            ["单箱约束违反", "−15 × 次数", "惩罚小箱残留"],
            ["IV 保留率", "+12 × min(保留率, 1.5)", "奖励区分度保留"],
            ["最小相邻风险差距", "+100 × max(0, 最小差距)", "奖励相邻档风险分离度"],
            ["档位距离", "−1.5 × |档位数 − 7|", "偏好目标档位数"],
            ["极端边界跨越", "−50 × 次数", "惩罚越过极端圈选边界"],
        ],
        widths=[5.0, 5.6, 5.4],
        caption_text="表 5　候选方案评分权重",
        footnote_lines=[
            "选择顺序：硬约束 → Train 主指标倒挂数 → 约束违反数 → Train 全指标倒挂数 → 综合得分 → IV 保留率 → 档位距离，取第一名。",
            "权重为经验值，只影响候选排序，不影响硬约束判定。",
        ],
    )

    # ============ 四、策略阈值设定 ============
    heading(doc, 1, "四、策略阈值设定")
    body(doc, "在完整 Train 上，沿低风险到高风险方向逐档放宽阈值（候选阈值为最终箱右边界，末箱用 Train 最大分数保留“全量通过”点），构造累计与边际风险曲线：累计指标反映放宽到该阈值的整体规模与风险；边际指标反映新增人群风险，用于识别风险拐点。")
    add_table(
        doc,
        ["约束阶段", "累计 1M30+ 笔数逾期率", "累计 3M30+ 笔数逾期率", "边际 3M30+ 笔数逾期率"],
        [
            ["自动通过", "≤ 0.90%", "≤ 5.50%", "≤ 9.00%"],
            ["总接纳（自动 + 人工）", "≤ 1.30%", "≤ 7.50%", "≤ 17.00%"],
        ],
        widths=[3.6, 4.2, 4.2, 4.0],
        caption_text="表 6　策略风险约束（示例配置）",
    )
    body(doc, "满足约束且累计通过率最高的阈值当选，形成三段策略：")
    formula(doc, "自动通过：score ≤ 自动通过阈值", "3")
    formula(doc, "人工审核：自动通过阈值 < score ≤ 总接纳阈值", "4")
    formula(doc, "拒绝：score > 总接纳阈值", "5")
    body(doc, "总接纳阈值不得严于自动通过阈值（若出现则对齐）；约束给出风险上限，目标为通过率最大化，实现“风险可控前提下业务量最大”的平衡。")

    # ============ 五、方案验证 ============
    heading(doc, 1, "五、方案验证")
    add_table(
        doc,
        ["验证维度", "指标", "口径 / 判定"],
        [
            ["风险排序", "单调性", "Train / OOT × 四类风险率是否非递减，输出倒挂次数与位置"],
            ["分布稳定", "PSI", "Train 与 OOT 各档样本占比差异；< 0.10 稳定，0.10~0.25 关注，≥ 0.25 明显变化"],
            ["区分能力", "AUC / KS", "Train / OOT × 1M30+ / 3M30+ 分别计算"],
            ["时间稳定", "月度稳定性", "逐月检查主指标单调性、倒挂次数、最大单次风险跌幅，识别异常月份"],
            ["策略效果", "分段验证", "自动通过 / 人工审核 / 拒绝三段在 Train 与 OOT 的规模与风险，应呈“低 < 中 < 高”梯度且不反转"],
        ],
        widths=[2.4, 2.8, 10.8],
        caption_text="表 7　验证指标体系",
    )
    body(doc, "验证分为两个层次：Train 检查方案拟合质量，OOT 作为独立数据进行最终样本外确认。")

    # ============ 六、局限性与待讨论问题 ============
    heading(doc, 1, "六、局限性与待讨论问题")
    body(doc, "以下局限按对上线决策的影响程度分为三级：高——上线前必须解决；中——上线前应明确处理方案；低——持续优化，不阻塞上线。")

    priority_label(doc, "高优先级 · 上线前必须解决")
    bullet(doc, "拒绝样本与选择偏差：三段阈值与约束仅基于已获贷人群（自动通过 + 人工审核通过）的表现，被拒人群无标签、不进入逾期率估计；若模型在被拒人群上排序失效，实际风险可能被系统性低估，未引入拒绝推断或补充样本校准。")
    bullet(doc, "尾部箱风险估计精度：坏样本量下限 20 仅保证点估计，最坏端箱的 3M30+ 率置信区间宽度可达 ±10~20 个百分点，策略边际风险线可能落入区间内而无法与相邻约束水平可靠区分；未引入置信区间或精确检验（如 Clopper-Pearson）作为达标判定与约束加严的依据。")
    bullet(doc, "目标函数与约束取值：阈值优化目标仅为累计通过率最大化，未嵌入逾期损失、资金成本与收入模型，也未量化收严 / 放松一档对规模与风险的边际影响；表 6 约束值为示例配置，其业务依据需评审确认。")
    bullet(doc, "阈值边界取整与上线执行细节（边界落入真实分数值）待细化。")

    priority_label(doc, "中优先级 · 上线前应明确处理方案")
    bullet(doc, "分箱保鲜期与刷新机制：边界与阈值绑定当前 score_mlt 版本，模型重构、重校准或客群结构变化后边界逐步退化；未定义定期重估周期、刷新触发条件及新旧分箱并轨验证（shadow）安排。")
    bullet(doc, "账龄结构干扰：1M30+ / 3M30+ 需要 1~3 个月观察窗，Train 尾部与 OOT 前期账龄未成熟、成熟率低，月度稳定性与分月指标易受账龄结构变化（而非真实风险变化）干扰；未按账龄分层复核成熟度充分性。")
    bullet(doc, "当前仅输出一套“平衡型”策略，未做保守 / 增长多策略对比与人工审核产能约束的联动分析。")
    bullet(doc, "金额口径依赖剩余本金估计字段，其估计准确性直接影响金额指标可信度。")
    bullet(doc, "候选评分权重与月度倒挂容忍度（0.3pp）等经验值，未做系统性敏感性分析。")

    priority_label(doc, "低优先级 · 持续优化")
    bullet(doc, "候选择优的抽样稳定性：候选评分与选择仅在 Train 上执行一次，未用 bootstrap / 子样本扰动验证最优方案与箱边界是否会翻转；最优与次优候选得分接近时结论稳健性未知。")
    bullet(doc, "初始 20 箱、最终 6~8 档（目标 7）为经验设定，未用数据驱动方式确定最优档位数。")
    bullet(doc, "单调性仅做相邻箱检查，无全局趋势检验（如 Cochran-Armitage）；金额口径单调性未纳入硬约束。")
    bullet(doc, "极端人群圈选数量（最好 / 最坏各 1 箱）为主观设定，其敏感性未检验。")

    doc.save(OUTPUT)
    print("generated:", OUTPUT)


if __name__ == "__main__":
    main()
